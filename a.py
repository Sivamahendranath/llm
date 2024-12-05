import streamlit as st
import os
import json
import traceback
from datetime import datetime
from bs4 import BeautifulSoup
import PyPDF2
from docx import Document
import requests
import pytesseract
from pdf2image import convert_from_path
from io import BytesIO
import google.generativeai as genai

# Streamlit Page Configuration
st.set_page_config(page_title="Entity Extraction Assistant", page_icon="🔍", layout="wide")

# API Configuration
try:
    genai.configure(api_key="AIzaSyBenjXh60oJf8B960QWrQxKAB6uYdxzDI8")  # Replace with your valid API key
    gemini = genai.GenerativeModel("gemini-pro")
except Exception as api_error:
    st.error("Error initializing Generative AI API. Check your API key.")
    st.stop()

# Utility Functions
def extract_text_from_pdf_with_ocr(pdf_file):
    """Extract text from image-based PDFs using OCR."""
    try:
        images = convert_from_path(pdf_file)
        text = ""
        for image in images:
            text += pytesseract.image_to_string(image)
        return text
    except Exception as e:
        st.error(f"OCR extraction failed: {e}")
        return ""

def extract_text_from_source(source, source_type):
    """Extract text from various sources."""
    try:
        if source_type == "URL":
            response = requests.get(source, timeout=10)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, "html.parser")
            return soup.get_text(strip=True)

        elif source_type == "PDF":
            if isinstance(source, BytesIO):  # Uploaded file
                pdf_reader = PyPDF2.PdfReader(source)
                text = " ".join(page.extract_text() or "" for page in pdf_reader.pages)
            else:  # File path
                with open(source, "rb") as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = " ".join(page.extract_text() or "" for page in pdf_reader.pages)

            if not text.strip():  # Fallback to OCR if no text is extracted
                st.warning("PDF contains no extractable text. Using OCR...")
                if isinstance(source, BytesIO):
                    temp_file_path = "temp_uploaded_pdf.pdf"
                    with open(temp_file_path, "wb") as temp_file:
                        temp_file.write(source.read())
                    text = extract_text_from_pdf_with_ocr(temp_file_path)
                else:
                    text = extract_text_from_pdf_with_ocr(source)
            return text

        elif source_type == "DOCX":
            if isinstance(source, BytesIO):  # Uploaded file
                doc = Document(source)
            else:  # File path
                doc = Document(source)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs)

        elif source_type == "TXT":
            if isinstance(source, BytesIO):  # Uploaded file
                return source.read().decode("utf-8")
            else:  # File path
                with open(source, "r", encoding="utf-8") as file:
                    return file.read()

        else:
            return ""
    except Exception as e:
        st.error(f"Error extracting text from {source_type}: {e}")
        st.error(traceback.format_exc())
        return ""

def chunk_text(text, chunk_size=500):
    """Split large text into manageable chunks."""
    words = text.split()
    return [" ".join(words[i:i + chunk_size]) for i in range(0, len(words), chunk_size)]

# Main Function
def main():
    st.title("🔍 Entity Extraction Assistant")
    st.sidebar.header("Configuration")

    # Input Options
    input_type = st.sidebar.selectbox("Input Type", ["Direct Text", "URL", "PDF", "DOCX", "TXT"])
    text_input = None

    # Input Handling
    if input_type == "Direct Text":
        text_input = st.text_area("Enter Text")
    elif input_type == "URL":
        text_input = st.text_input("Enter URL")
    elif input_type in ["PDF", "DOCX", "TXT"]:
        text_input = st.file_uploader(f"Upload {input_type} File", type=input_type.lower())

    # Extract Button
    if st.sidebar.button("Extract Entities"):
        if not text_input:
            st.warning("Please provide input data.")
            return

        st.subheader("Processing...")
        try:
            # Extract text
            if input_type in ["PDF", "DOCX", "TXT"] and isinstance(text_input, BytesIO):
                extracted_text = extract_text_from_source(text_input, input_type)
            elif input_type == "URL":
                extracted_text = extract_text_from_source(text_input, "URL")
            else:
                extracted_text = text_input

            if not extracted_text.strip():
                st.error("No text extracted from the input.")
                return

            # Display Extracted Text
            st.subheader("Extracted Text")
            st.text(extracted_text[:1000])  # Show first 1000 characters for verification

            # Simulated Entity Extraction (Replace with actual Gemini integration)
            st.success("Entities Extracted Successfully!")
            st.json([{"name": "John Doe", "category": "PERSON", "context": "Example context"}])

        except Exception as e:
            st.error(f"Unexpected error occurred: {e}")
            st.error(traceback.format_exc())

if __name__ == "__main__":
    main()
